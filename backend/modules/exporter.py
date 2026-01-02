"""
Exporter Module
Exports questions and results to PDF and DOCX formats
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from typing import List, Dict, Optional


class ResultExporter:
    """Exports quiz questions and results to various formats"""

    def __init__(self):
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)

    def export(
        self,
        questions: List[Dict],
        responses: Optional[List[Dict]],
        export_type: str,
        file_format: str,
        session_id: str
    ) -> str:
        """
        Export questions or results

        Args:
            questions: List of questions
            responses: List of responses (None for questions_only)
            export_type: "questions_only" or "results_with_answers"
            file_format: "pdf" or "docx"
            session_id: Quiz session ID

        Returns:
            Path to exported file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{export_type}_{session_id}_{timestamp}"

        if file_format == "pdf":
            return self._export_to_pdf(
                questions, responses, export_type, filename
            )
        elif file_format == "docx":
            return self._export_to_docx(
                questions, responses, export_type, filename
            )
        else:
            raise ValueError(f"Unsupported format: {file_format}")

    def _export_to_pdf(
        self,
        questions: List[Dict],
        responses: Optional[List[Dict]],
        export_type: str,
        filename: str
    ) -> str:
        """Export to PDF format"""
        filepath = os.path.join(self.export_dir, f"{filename}.pdf")

        # Create PDF
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(0, 0, 139),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=RGBColor(0, 0, 0),
            spaceAfter=12
        )

        # Title
        title = "Quiz Questions" if export_type == "questions_only" else "Quiz Results"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3 * inch))

        # Metadata
        date_str = datetime.now().strftime("%B %d, %Y at %I:%M %p")
        story.append(Paragraph(f"Generated on: {date_str}", styles['Normal']))
        story.append(Paragraph(f"Total Questions: {len(questions)}", styles['Normal']))
        story.append(Spacer(1, 0.3 * inch))

        # Questions
        for idx, question in enumerate(questions, 1):
            # Question number and text
            story.append(Paragraph(
                f"<b>Question {idx}:</b> {question.get('question_text', '')}",
                heading_style
            ))
            story.append(Spacer(1, 0.1 * inch))

            # Options
            for option in ['A', 'B', 'C', 'D']:
                option_key = f"option_{option.lower()}"
                option_text = question.get(option_key, '')

                # Highlight correct answer if showing results
                if export_type == "results_with_answers":
                    if option == question.get('correct_answer'):
                        story.append(Paragraph(
                            f"<b>{option}. {option_text} ✓</b>",
                            styles['Normal']
                        ))
                    else:
                        story.append(Paragraph(
                            f"{option}. {option_text}",
                            styles['Normal']
                        ))
                else:
                    story.append(Paragraph(
                        f"{option}. {option_text}",
                        styles['Normal']
                    ))

            story.append(Spacer(1, 0.1 * inch))

            # Show explanation if results mode
            if export_type == "results_with_answers":
                explanation = question.get('explanation', '')
                if explanation:
                    story.append(Paragraph(
                        f"<i>Explanation: {explanation}</i>",
                        styles['Normal']
                    ))
                    story.append(Spacer(1, 0.1 * inch))

                # Show user answer if responses provided
                if responses:
                    user_resp = next(
                        (r for r in responses if r.get('question_id') == question.get('id')),
                        None
                    )
                    if user_resp:
                        user_ans = user_resp.get('user_answer', 'Not answered')
                        is_correct = user_resp.get('is_correct', False)
                        status = "✓ Correct" if is_correct else "✗ Incorrect"

                        story.append(Paragraph(
                            f"Your Answer: {user_ans} - {status}",
                            styles['Normal']
                        ))

            story.append(Spacer(1, 0.3 * inch))

            # Page break every 5 questions
            if idx % 5 == 0 and idx < len(questions):
                story.append(PageBreak())

        # Build PDF
        doc.build(story)
        return filepath

    def _export_to_docx(
        self,
        questions: List[Dict],
        responses: Optional[List[Dict]],
        export_type: str,
        filename: str
    ) -> str:
        """Export to DOCX format"""
        filepath = os.path.join(self.export_dir, f"{filename}.docx")

        # Create document
        doc = Document()

        # Title
        title = "Quiz Questions" if export_type == "questions_only" else "Quiz Results"
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Total Questions: {len(questions)}")
        doc.add_paragraph()

        # Questions
        for idx, question in enumerate(questions, 1):
            # Question heading
            q_heading = doc.add_heading(f"Question {idx}", level=2)

            # Question text
            q_text = doc.add_paragraph(question.get('question_text', ''))
            q_text.runs[0].bold = True

            # Options
            for option in ['A', 'B', 'C', 'D']:
                option_key = f"option_{option.lower()}"
                option_text = question.get(option_key, '')

                opt_para = doc.add_paragraph(f"{option}. {option_text}")

                # Highlight correct answer
                if export_type == "results_with_answers":
                    if option == question.get('correct_answer'):
                        opt_para.runs[0].bold = True
                        opt_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
                        opt_para.add_run(" ✓").bold = True

            # Explanation
            if export_type == "results_with_answers":
                explanation = question.get('explanation', '')
                if explanation:
                    exp_para = doc.add_paragraph(f"Explanation: {explanation}")
                    exp_para.runs[0].italic = True
                    exp_para.runs[0].font.size = Pt(10)

                # User response
                if responses:
                    user_resp = next(
                        (r for r in responses if r.get('question_id') == question.get('id')),
                        None
                    )
                    if user_resp:
                        user_ans = user_resp.get('user_answer', 'Not answered')
                        is_correct = user_resp.get('is_correct', False)

                        ans_para = doc.add_paragraph(f"Your Answer: {user_ans}")

                        if is_correct:
                            ans_para.add_run(" - ✓ Correct").bold = True
                            ans_para.runs[-1].font.color.rgb = RGBColor(0, 128, 0)
                        else:
                            ans_para.add_run(" - ✗ Incorrect").bold = True
                            ans_para.runs[-1].font.color.rgb = RGBColor(255, 0, 0)

            doc.add_paragraph()  # Space between questions

        # Save document
        doc.save(filepath)
        return filepath
